from io import BytesIO
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont  # Добавляем поддержку TTF шрифтов
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

# Регистрируем шрифты для поддержки кириллицы
def register_fonts():
    """Регистрация шрифтов для поддержки кириллицы"""
    try:
        # Попробуем найти стандартные шрифты с кириллицей
        font_paths = [
            'C:/Windows/Fonts/arial.ttf',  # Windows
            'C:/Windows/Fonts/times.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',  # Linux
            '/System/Library/Fonts/Arial.ttf',  # macOS
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    # Регистрируем обычный и жирный шрифт
                    pdfmetrics.registerFont(TTFont('Arial', font_path))
                    pdfmetrics.registerFont(TTFont('Arial-Bold', font_path))
                    return True
                except:
                    continue
        
        # Если стандартные шрифты не найдены, используем встроенные
        # ReportLab имеет базовую поддержку кириллицы через шрифт Helvetica
        return False
    except:
        return False

# Регистрируем шрифты при импорте модуля
fonts_registered = register_fonts()

def export_to_pdf(reports_data, report_type, report_name):
    """Экспорт отчета в PDF"""
    try:
        buffer = BytesIO()
        
        # Создаем документ с указанием шрифта по умолчанию
        if fonts_registered:
            doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=72,
                                  fontName='Arial')
        else:
            doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=72)
        
        elements = []
        
        styles = getSampleStyleSheet()
        
        # Создаем кастомные стили с учетом кириллицы
        if fonts_registered:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName='Arial-Bold',
                fontSize=16,
                spaceAfter=30,
                alignment=1,  # Center
                textColor=colors.HexColor('#2C3E50')
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontName='Arial-Bold',
                fontSize=12,
                spaceAfter=20,
                textColor=colors.HexColor('#34495E')
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName='Arial',
                fontSize=10,
                spaceAfter=10
            )
        else:
            # Используем стандартные шрифты если Arial не доступен
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1,
                textColor=colors.HexColor('#2C3E50')
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=20,
                textColor=colors.HexColor('#34495E')
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=10
            )
        
        # Заголовок отчета
        title = f"ОТЧЕТ: {get_report_title(report_type, report_name)}"
        elements.append(Paragraph(title, title_style))
        
        # Информация о генерации
        elements.append(Paragraph(f"<b>Дата генерации:</b> {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}", normal_style))
        elements.append(Paragraph(f"<b>Тип отчета:</b> {report_type.upper()}", normal_style))
        elements.append(Spacer(1, 20))
        
        # Данные отчета
        if report_type == 'courier':
            data_table = prepare_courier_table(reports_data, report_name)
        else:
            data_table = prepare_course_table(reports_data, report_name)
        
        if data_table:
            # Создаем таблицу с данными
            table = Table(data_table, colWidths=[doc.width/len(data_table[0])] * len(data_table[0]))
            
            # Стили для таблицы
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold' if not fonts_registered else 'Arial-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#DEE2E6')),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')]),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica' if not fonts_registered else 'Arial'),
            ])
            
            table.setStyle(table_style)
            elements.append(table)
            elements.append(Spacer(1, 30))
            
            # Статистика
            elements.append(Paragraph("<b>СТАТИСТИКА</b>", subtitle_style))
            stats = get_report_statistics(reports_data, report_type, report_name)
            for stat in stats:
                # Экранируем символы для корректного отображения
                stat_escaped = stat.replace('°', 'градусов').replace('±', '+-')
                elements.append(Paragraph(f"• {stat_escaped}", normal_style))
            
            # Подпись
            elements.append(Spacer(1, 50))
            elements.append(Paragraph("___________________________", normal_style))
            elements.append(Paragraph("<i>Генератор отчетов</i>", normal_style))
        else:
            elements.append(Paragraph("<b>Нет данных для отображения</b>", normal_style))
        
        try:
            doc.build(elements)
        except Exception as build_error:
            print(f"Ошибка при сборке PDF: {build_error}")
            # Попробуем альтернативный метод с простым текстом
            return export_to_pdf_simple(reports_data, report_type, report_name)
        
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"Ошибка при генерации PDF: {e}")
        return None

def export_to_pdf_simple(reports_data, report_type, report_name):
    """Простой экспорт PDF (запасной вариант)"""
    try:
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        
        # Устанавливаем шрифт
        c.setFont("Helvetica", 12)
        
        # Заголовок
        title = f"ОТЧЕТ: {get_report_title(report_type, report_name)}"
        c.drawString(100, 800, title)
        
        # Дата
        c.setFont("Helvetica", 10)
        c.drawString(100, 780, f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
        c.drawString(100, 765, f"Тип отчета: {report_type.upper()}")
        
        # Данные
        y_position = 730
        if report_type == 'courier':
            data = prepare_courier_table_simple(reports_data, report_name)
        else:
            data = prepare_course_table_simple(reports_data, report_name)
        
        if data:
            c.setFont("Helvetica-Bold", 10)
            # Заголовки таблицы
            for i, header in enumerate(data[0]):
                c.drawString(100 + i * 100, y_position, str(header)[:15])
            
            y_position -= 20
            c.setFont("Helvetica", 9)
            
            # Данные (ограничиваем количество строк)
            for row in data[1:20]:  # Максимум 20 строк
                if y_position < 50:
                    c.showPage()
                    y_position = 750
                    c.setFont("Helvetica", 9)
                
                for i, cell in enumerate(row):
                    cell_text = str(cell)[:15]  # Обрезаем длинный текст
                    c.drawString(100 + i * 100, y_position, cell_text)
                y_position -= 15
        
        c.save()
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"Ошибка при генерации простого PDF: {e}")
        return None

def prepare_courier_table_simple(reports_data, report_name):
    """Простая подготовка таблицы для курьерской доставки"""
    if report_name not in reports_data['courier_reports']:
        return None
    
    data = reports_data['courier_reports'][report_name]
    
    if not data:
        return [['Нет данных']]
    
    if report_name == 'courier_stats':
        return [
            ['Курьер', 'Кол-во', 'Вес (кг)'],
            *[[str(item.get('_id', 'Не указан'))[:10], 
               str(item.get('count', 0)),
               f"{item.get('total_weight', 0):.1f}"] for item in data[:10]]
        ]
    elif report_name == 'heavy_parcels':
        return [
            ['Отправитель', 'Получатель', 'Вес'],
            *[[str(item['sender']['full_name'])[:10],
               str(item['receiver']['full_name'])[:10],
               f"{item['parcel']['weight']:.1f}"] for item in data[:10]]
        ]
    
    return [['Данные недоступны']]

def prepare_course_table_simple(reports_data, report_name):
    """Простая подготовка таблицы для курсов"""
    if report_name not in reports_data['courses_reports']:
        return None
    
    data = reports_data['courses_reports'][report_name]
    
    if not data:
        return [['Нет данных']]
    
    if report_name == 'department_stats':
        return [
            ['Отдел', 'Сотрудники', 'Курсы'],
            *[[str(item.get('department', 'Не указан'))[:10],
               str(item.get('employee_count', 0)),
               str(item.get('course_count', 0))] for item in data[:10]]
        ]
    elif report_name in ['upcoming_courses', 'long_courses']:
        return [
            ['Курс', 'Преподаватель', 'Часы'],
            *[[str(item['course_name'])[:10],
               str(item['teacher']['name'])[:10],
               str(item['hours'])] for item in data[:10]]
        ]
    
    return [['Данные недоступны']]

# Остальные функции оставляем без изменений...
def export_to_docx(reports_data, report_type, report_name):
    """Экспорт отчета в DOCX"""
    try:
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # Заголовок
        title = doc.add_heading(f'ОТЧЕТ: {get_report_title(report_type, report_name)}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о отчете
        doc.add_paragraph(f'Дата генерации: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
        doc.add_paragraph(f'Тип отчета: {report_type.upper()}')
        doc.add_paragraph()
        
        # Данные отчета
        if report_type == 'courier':
            data = prepare_courier_table(reports_data, report_name)
        else:
            data = prepare_course_table(reports_data, report_name)
        
        if data:
            # Создаем таблицу
            table = doc.add_table(rows=1, cols=len(data[0]))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(data[0]):
                hdr_cells[i].text = str(header)
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Данные таблицы
            for row_data in data[1:]:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = str(cell_data)
            
            doc.add_paragraph()
            
            # Статистика
            stats_heading = doc.add_heading('Статистика', 2)
            stats = get_report_statistics(reports_data, report_type, report_name)
            for stat in stats:
                doc.add_paragraph(f'• {stat}', style='List Bullet')
            
            # Подпись
            doc.add_paragraph()
            doc.add_paragraph('_' * 40)
            doc.add_paragraph('Генератор отчетов', style='Intense Quote')
        
        # Сохраняем в буфер
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"Ошибка при генерации DOCX: {e}")
        return None

def get_report_title(report_type, report_name):
    """Получение заголовка отчета"""
    titles = {
        'courier': {
            'heavy_parcels': 'Тяжелые посылки (>5 кг)',
            'in_transit': 'Посылки в пути',
            'last_week': 'Посылки за последнюю неделю',
            'by_sender': 'Посылки по отправителям',
            'courier_stats': 'Статистика по курьерам',
            'all': 'Все посылки'
        },
        'courses': {
            'upcoming_courses': 'Предстоящие курсы',
            'long_courses': 'Длительные курсы (>40 часов)',
            'by_teacher': 'Курсы по преподавателям',
            'full_courses': 'Курсы с полными группами',
            'department_stats': 'Статистика по отделам',
            'all': 'Все курсы'
        }
    }
    return titles.get(report_type, {}).get(report_name, 'Общий отчет')

def prepare_courier_table(reports_data, report_name):
    """Подготовка таблицы данных для курьерской доставки"""
    if report_name not in reports_data['courier_reports']:
        return None
    
    data = reports_data['courier_reports'][report_name]
    
    if not data:
        return [['Нет данных для отображения']]
    
    if report_name == 'courier_stats':
        return [
            ['Курьер', 'Количество посылок', 'Общий вес (кг)', 'Средний вес'],
            *[[str(item.get('_id', 'Не указан'))[:20], 
               str(item.get('count', 0)), 
               f"{item.get('total_weight', 0):.2f}",
               f"{item.get('total_weight', 0)/item.get('count', 1):.2f}" if item.get('count', 0) > 0 else "0.00"] 
              for item in data]
        ]
    elif report_name == 'heavy_parcels':
        return [
            ['Трек №', 'Отправитель', 'Получатель', 'Вес (кг)', 'Статус', 'Дата отправки'],
            *[[item.get('tracking_number', 'N/A'),
               str(item['sender']['full_name'])[:15],
               str(item['receiver']['full_name'])[:15],
               f"{item['parcel']['weight']:.2f}",
               item['status'],
               item['dates']['dispatch_date']] for item in data[:20]]
        ]
    elif report_name == 'in_transit':
        return [
            ['Трек №', 'Отправитель', 'Получатель', 'Курьер', 'Ожидаемая дата', 'Стоимость'],
            *[[item.get('tracking_number', 'N/A'),
               str(item['sender']['full_name'])[:12],
               str(item['receiver']['full_name'])[:12],
               str(item['courier']['name'])[:12],
               item['dates']['delivery_date'],
               f"{item.get('delivery_cost', 0):.2f} руб."] for item in data[:20]]
        ]
    elif report_name == 'last_week':
        return [
            ['Трек №', 'Отправитель', 'Статус', 'Дата отправки', 'Дата получения', 'Вес (кг)'],
            *[[item.get('tracking_number', 'N/A'),
               str(item['sender']['full_name'])[:15],
               item['status'],
               item['dates']['dispatch_date'],
               item['dates'].get('actual_delivery_date', 'Не доставлено'),
               f"{item['parcel']['weight']:.2f}"] for item in data[:20]]
        ]
    
    return None

def prepare_course_table(reports_data, report_name):
    """Подготовка таблицы данных для курсов"""
    if report_name not in reports_data['courses_reports']:
        return None
    
    data = reports_data['courses_reports'][report_name]
    
    if not data:
        return [['Нет данных для отображения']]
    
    if report_name == 'department_stats':
        return [
            ['Отдел', 'Количество сотрудников', 'Количество курсов', 'Среднее кол-во сотрудников'],
            *[[str(item.get('department', 'Не указан'))[:20], 
               str(item.get('employee_count', 0)), 
               str(item.get('course_count', 0)),
               f"{item.get('employee_count', 0)/item.get('course_count', 1):.1f}" if item.get('course_count', 0) > 0 else "0.0"] 
              for item in data]
        ]
    elif report_name in ['upcoming_courses', 'long_courses', 'full_courses']:
        return [
            ['Код курса', 'Название курса', 'Преподаватель', 'Даты', 'Часы', 'Стоимость', 'Участников'],
            *[[item.get('course_code', 'N/A'),
               str(item['course_name'])[:20],
               str(item['teacher']['name'])[:15],
               f"{item['dates']['start_date']} - {item['dates']['end_date']}",
               str(item['hours']),
               f"{item.get('price', 0):.2f} руб.",
               str(len(item.get('employees', [])))] for item in data[:20]]
        ]
    
    return None

def get_report_statistics(reports_data, report_type, report_name):
    """Получение статистики по отчету"""
    stats = []
    
    if report_type == 'courier':
        data = reports_data['courier_reports'].get(report_name, [])
        if report_name == 'courier_stats':
            total_parcels = sum(item.get('count', 0) for item in data)
            total_weight = sum(item.get('total_weight', 0) for item in data)
            avg_weight = total_weight / total_parcels if total_parcels > 0 else 0
            stats.extend([
                f"Всего курьеров: {len(data)}",
                f"Общее количество посылок: {total_parcels}",
                f"Общий вес всех посылок: {total_weight:.2f} кг",
                f"Средний вес посылки: {avg_weight:.2f} кг"
            ])
        else:
            stats.append(f"Количество записей: {len(data)}")
            if data:
                total_weight = sum(item.get('parcel', {}).get('weight', 0) for item in data)
                stats.append(f"Общий вес: {total_weight:.2f} кг")
    
    elif report_type == 'courses':
        data = reports_data['courses_reports'].get(report_name, [])
        if report_name == 'department_stats':
            total_employees = sum(item.get('employee_count', 0) for item in data)
            total_courses = sum(item.get('course_count', 0) for item in data)
            avg_employees = total_employees / len(data) if data else 0
            stats.extend([
                f"Всего отделов: {len(data)}",
                f"Общее количество сотрудников: {total_employees}",
                f"Общее количество курсов: {total_courses}",
                f"Среднее количество сотрудников на отдел: {avg_employees:.1f}"
            ])
        else:
            stats.append(f"Количество курсов: {len(data)}")
            if data:
                total_hours = sum(item.get('hours', 0) for item in data)
                total_participants = sum(len(item.get('employees', [])) for item in data)
                total_price = sum(item.get('price', 0) for item in data)
                stats.extend([
                    f"Общее количество часов: {total_hours}",
                    f"Общее количество участников: {total_participants}",
                    f"Общая стоимость всех курсов: {total_price:.2f} руб.",
                    f"Средняя стоимость курса: {total_price/len(data):.2f} руб." if data else "0.00 руб."
                ])
    
    return stats